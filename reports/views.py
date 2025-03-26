import io
import datetime
from django.http import JsonResponse
import requests
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView,FormView
from django.contrib.auth.mixins import LoginRequiredMixin
from docx import Document
from django.conf import settings
from .models import Patient, Provider, Report
from .forms import CreateUserForm, PatientForm, ProviderForm, ReportGenerationForm
from django.shortcuts import render
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.core.files import File
from django.db.models import Q

@login_required
@user_passes_test(lambda u: u.is_superuser)
def admin_panel(request):
    users = User.objects.all()
    return render(request, 'admin_panel.html', {'users': users})

class CreateUserView(LoginRequiredMixin, CreateView):
    model = User
    template_name = 'create_user.html'
    form_class = CreateUserForm
    success_url = reverse_lazy('admin_panel')

class UserDeleteView(LoginRequiredMixin, DeleteView):
    
    model = User
    template_name = 'confirm_delete.html'
    success_url = reverse_lazy('admin_panel')



class GenerateReportsView(LoginRequiredMixin, FormView):
    model = Report
    form_class = ReportGenerationForm
    template_name = 'generate_reports.html'
    success_url = reverse_lazy('report_list')  # Redirects to a page listing reports after generation

    def get_ai_response(self, client_name, client_support_plan):
        """
        Calls the Ollama API to generate an AI response based on patient name and support plan.
        Matches the prompt structure from the desktop application's AIResponse function.
        """
        prompt = (
            f"You are a care provider. Your patient is {client_name}. They are on the {client_support_plan} support plan. "
            "You are required to provide service based on the agency for persons with disabilities person-centered approach. "
            "What tasks did you assist with today? Limit your response to three paragraphs. "
            "How did they react to you assisting them? Limit your response to two paragraphs. "
            "Give me a problem, action, assistance, and solution (positive or negative) for today. "
            "What is a random question you asked the patient and their response? Limit your response to one sentence."
        )
        try:
            response = requests.post(
                f'{settings.AI_SERVICE_URL}/api/chat',  # Assumes Ollama runs locally; adjust for production
                json={
                    'model': 'llama3.1',  # Matches the model used in the desktop app
                    'messages': [{'role': 'user', 'content': prompt}],
                    "stream": False
                }
            )
            if response.status_code == 200:
                return response.json()['message']['content']
            else:
                raise Exception(f"AI API request failed with status code {response.status_code}")
        except Exception as e:
            # In a production environment, consider logging the error instead of raising it
            raise Exception(f"Failed to get AI response: {str(e)}")

    def form_valid(self, form):
        """
        Handles the form submission by generating reports for each day in the date range.
        Creates a DOCX file for each report and saves it to the Report model.
        """
        # Extract form data
        patient = form.cleaned_data['patient']
        provider = form.cleaned_data['provider']
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        start_time = form.cleaned_data['start_time']
        end_time = form.cleaned_data['end_time']

        # Calculate total time (assumed constant across all days, as in the desktop app)
        start_datetime = datetime.datetime.combine(start_date, start_time)
        end_datetime = datetime.datetime.combine(start_date, end_time)
        total_time = end_datetime - start_datetime
        total_hours = total_time.total_seconds() / 3600  # Convert to hours
        time_quarter = total_hours * 4  # Quarter hours, matching desktop app logic

        # Generate reports for each day in the date range
        current_date = start_date
        while current_date <= end_date:
            # Generate AI response for the current day
            ai_response = self.get_ai_response(patient.name, patient.support_plan)
            
            # Create a new DOCX document
            doc = Document()

            # Header section (mimics pDocx function from desktop app)
            p = doc.add_paragraph(f"Service Log\n{current_date.strftime('%m-%d-%Y')}\n")
            p.add_run(f"{patient.name}\n")
            p.add_run(f"Medicaid Number: {patient.id_number}\n")  # Assumes Patient model has id_number
            p.add_run(f"{patient.service}\n")  # Assumes Patient model has service
            p.add_run(f"{provider.name}\n")
            p.add_run("Peterson Family Care LLC Provider Number: 009279700\n")
            p.add_run(f"Start Time: {start_time.strftime('%I:%M %p')} - End Time: {end_time.strftime('%I:%M %p')}\n")
            p.add_run(f"Total Hours: {total_hours:.2f}\n")
            p.add_run(f"Total Quarter Hours: {time_quarter:.2f}\n")

            # AI-generated body section
            doc.add_heading('Daily Report', level=1)
            doc.add_paragraph(ai_response)

            # Signature section (simplified as placeholders)
            signature = doc.add_paragraph()
            try:
                names = provider.name.split(' ')
                sig = ''
                for name in names:
                    sig = sig + name[0]
                signature.add_run(sig + '\n').font.name='Segoe Script'
            except:
                signature.add_run(patient.name + '\n').font.name='Segoe Script'
            signature.add_run(provider.name + '\n').font.name='Segoe Script'

            # Save DOCX to an in-memory BytesIO object
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)

            # Create and save Report instance
            report = Report(
                patient=patient,
                provider=provider,
                date=current_date,
                start_time=start_time,
                end_time=end_time,
                file=File(
                    docx_io,
                    name=f"{patient.name.replace(' ', '-')}_{current_date.strftime('%m-%d-%Y')}.docx"
                )
            )
            report.save()

            # Move to the next day
            current_date += datetime.timedelta(days=1)

        # Redirect to success URL after all reports are generated
        return super().form_valid(form)

    def post(self, request, *args, **kwargs):
        """
        Overrides the default POST handler to process the form submission.
        The original print(request.POST) is removed as it's no longer needed.
        """
        return super().post(request, *args, **kwargs)
    

# Patient Views
class PatientListView(LoginRequiredMixin, ListView):
    model = Patient
    template_name = 'patient_list.html'

class PatientCreateView(LoginRequiredMixin, CreateView):
    model = Patient
    form_class = PatientForm
    template_name = 'patient_form.html'
    success_url = reverse_lazy('patient_list')

class PatientUpdateView(LoginRequiredMixin, UpdateView):
    model = Patient
    form_class = PatientForm
    template_name = 'patient_form.html'
    success_url = reverse_lazy('patient_list')

class PatientDeleteView(LoginRequiredMixin, DeleteView):
    model = Patient
    template_name = 'patient_confirm_delete.html'
    success_url = reverse_lazy('patient_list')

# Provider Views
class ProviderListView(LoginRequiredMixin, ListView):
    model = Provider
    template_name = 'provider_list.html'

class ProviderCreateView(LoginRequiredMixin, CreateView):
    model = Provider
    form_class = ProviderForm
    template_name = 'provider_form.html'
    success_url = reverse_lazy('provider_list')

class ProviderUpdateView(LoginRequiredMixin, UpdateView):
    model = Provider
    form_class = ProviderForm
    template_name = 'provider_form.html'
    success_url = reverse_lazy('provider_list')

class ProviderDeleteView(LoginRequiredMixin, DeleteView):
    model = Provider
    template_name = 'provider_confirm_delete.html'
    success_url = reverse_lazy('provider_list')
    
class ReportListView(LoginRequiredMixin, ListView):
    model = Report
    template_name = 'report_list.html'
    ordering = ['patient__name', '-date']

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(patient__name__icontains=query) |
                Q(date__icontains=query)
            )
        return queryset
    
def get_report_preview(request, report_id):
    try:
        report = Report.objects.get(id=report_id)
        preview_text = report.get_preview_text()
        return JsonResponse({'preview': preview_text})
    except Report.DoesNotExist:
        return JsonResponse({'error': 'Report not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

class ReportDeleteView(LoginRequiredMixin, DeleteView):
    model = Report
    template_name = 'report_confirm_delete.html'
    success_url = reverse_lazy('report_list')