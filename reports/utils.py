# reports/utils.py
from ollama import chat
from docx import Document
from io import BytesIO

def AIResponse(ClientName, ClientSupportPlan, date):
    response = chat(model='llama3.1', messages=[
        {
            'role': 'user',
            'content': (
                f'You are a care provider. Your patient is {ClientName}. They are on the {ClientSupportPlan} '
                f'support plan. Today is {date.strftime("%B %d, %Y")}. What tasks did you assist with today? '
                'Limit your response to three paragraphs. How did they react to you assisting them? '
                'Limit your response to two paragraphs. Give me a problem, action, assistance, and solution '
                '(positive or negative) for today. What is a random question you asked the patient and their response? '
                'Limit your response to one sentence'
            ),
        }
    ])
    return response

def generate_report_docx(ClientName, ClientID, current_date, response, serviceProvided, serviceProvidedBy, startTime, endTime, totaltime, timeQuarter):
    document = Document()
    # Header
    p = document.add_paragraph(f'Service Log\n{current_date.strftime("%m-%d-%Y")}\n')
    p.add_run(f'{ClientName}\n')
    p.add_run(f'Medicaid Number: {ClientID}\n')
    p.add_run(f'{serviceProvided}\n')
    p.add_run(f'{serviceProvidedBy}\n')
    p.add_run('Peterson Family Care LLC Provider Number: 009279700\n')
    p.add_run(f'Start Time: {startTime} - End Time: {endTime}\n')
    p.add_run(f'Total Hours: {totaltime}\n')
    p.add_run(f'Total Quarter Hours: {timeQuarter}\n')
    # Body
    document.add_heading('Daily Report', level=1)
    document.add_paragraph(f'{response}\n')
    # Signatures
    signature = document.add_paragraph('Client Signature: ')
    names = ClientName.split(' ')
    sig = ''.join(name[0] for name in names)
    signature.add_run(f'{sig}\n').font.name = 'Segoe Script'
    signature.add_run('Provider Signature: ')
    signature.add_run(serviceProvidedBy).font.name = 'Segoe Script'
    # Save to BytesIO
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer